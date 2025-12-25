import datetime
from copy import deepcopy
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from modules.fetch_content import fetch_content

# ✅ 來源對應字典
SOURCE_MAP = {
    "nextapple.com": "壹蘋網",
    "ctinews.com": "中天網",
    "ctitv.com.tw": "中天網",
    "cti.com.tw": "中天網",
    "knews.com.tw": "知新聞",
    "ebc.net.tw": "東森網",
    "ctwant.com": "周刊王",
    "setn.com": "三立網",
    "ettoday.net": "東森雲",
    "udn.com": "聯合新聞網",
    "chinatimes.com": "中時新聞網",
    "mirrordaily.news": "鏡報",
    "tvbs.com.tw": "TVBS",
    "mirrormedia.mg": "鏡週刊",
    "mnews.tw": "鏡新聞",
    "ltn.com.tw": "自由時報",
    "cna.com.tw": "中央社",
}

def _source_from_url(url: str) -> str:
    """依網址網域判斷報別"""
    host = urlparse(url).hostname or ""
    host = host.lower()
    for key, source in SOURCE_MAP.items():
        if key in host:
            return source
    return host  # 預設用網域當來源

def _set_cell_style(cell, text, center=True, font_name="標楷體", font_size=14):
    """設定儲存格文字樣式"""
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.name = font_name or "Arial"  # ✅ fallback
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name or "Arial")
            run.font.size = Pt(font_size)

async def export_to_word_from_urls(urls, filename="新聞剪報.docx"):
    """接收 URL 清單，抓取新聞並匯出 Word 檔（async 版）"""
    try:
        template = Document("templates/新聞輸出範本.docx")
        doc = Document("templates/新聞輸出範本.docx")
    except FileNotFoundError as e:
        raise RuntimeError(f"範本檔不存在：{e}")
    except Exception as e:
        raise RuntimeError(f"無法開啟範本檔：{e}")

    base_title_el = template.paragraphs[0]._element
    base_table_el = template.tables[0]._element

    now = datetime.datetime.now()
    roc_year = now.year - 1911
    roc_date = f"{roc_year}-{now.strftime('%m-%d')}"  # ✅ 改用 strftime

    for idx, url in enumerate(urls, start=1):
        source = _source_from_url(url)

        if idx == 1:
            target_table = doc.tables[0]
        else:
            new_title_el = deepcopy(base_title_el)
            new_tbl_el = deepcopy(base_table_el)
            doc.element.body.append(new_title_el)
            doc.element.body.append(new_tbl_el)
            target_table = doc.tables[-1]

        # 填入表格欄位
        _set_cell_style(target_table.cell(0, 1), roc_date)      # 日期
        _set_cell_style(target_table.cell(0, 3), source)        # 報別
        _set_cell_style(target_table.cell(0, 5), str(idx))      # 頁碼
        _set_cell_style(target_table.cell(1, 1), "社會")        # 版別（暫時固定）

        try:
            # ✅ async 呼叫 fetch_content
            content = await fetch_content(url)
        except Exception as e:
            content = f"（抓取失敗: {e}）"

        body_cell = target_table.cell(2, 0)
        body_cell.text = content

        if template.tables[0].cell(2, 0).paragraphs:
            body_cell.paragraphs[0].style = template.tables[0].cell(2, 0).paragraphs[0].style

    doc.save(filename)
    return filename

# ✅ 提供別名，讓 Bot 可以用 export_to_word
export_to_word = export_to_word_from_urls
